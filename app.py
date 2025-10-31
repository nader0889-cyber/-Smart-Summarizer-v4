# app.py — Smart Summarizer v5 (Gemini + Supabase)
import streamlit as st
import google.generativeai as genai
from langdetect import detect
from docx import Document
import PyPDF2
import io, json, re, datetime
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from supabase import create_client

# -------------------- Page Config --------------------
st.set_page_config(page_title="Smart Summarizer v5", page_icon="🧠", layout="wide")

# -------------------- CSS & Background --------------------
st.markdown("""
<style>
:root {--g1:#0f172a; --g2:#021124; --g3:#062b2f;}
.app-bg { position: fixed; inset:0; background: linear-gradient(120deg,#0f172a,#0b3a4a,#09304b,#05323a);
background-size:400% 400%; animation: gradientMove 18s ease infinite; z-index:-1; filter: blur(20px) saturate(110%); opacity:0.95;}
@keyframes gradientMove {0%{background-position:0% 50%}50%{background-position:100% 50%}100%{background-position:0% 50%}}
textarea, .stTextArea textarea { background: rgba(0,0,0,0.2) !important; color: #f8f8f8 !important; border-radius: 8px !important; padding: 10px !important;}
.stButton>button { background: linear-gradient(90deg,#6ee7b7,#3b82f6) !important; color: #041025 !important; font-weight:700; border-radius:10px !important; padding: 8px 14px !important;}
</style><div class="app-bg"></div>
""", unsafe_allow_html=True)

st.title("✨ Smart Summarizer v5 — Gemini + Supabase")
st.markdown("**تلخيص ذكي، ترجمة، رفع ملفات، تخزين النتائج في Supabase، تنزيل ملفات**")

# -------------------- API Keys --------------------
if "GEMINI_API_KEY" not in st.secrets or "SUPABASE_URL" not in st.secrets or "SUPABASE_KEY" not in st.secrets:
    st.error("ضع GEMINI_API_KEY و SUPABASE_URL و SUPABASE_KEY في .streamlit/secrets.toml")
    st.stop()

genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
supabase = create_client(st.secrets["SUPABASE_URL"], st.secrets["SUPABASE_KEY"])

# -------------------- Helpers --------------------
def extract_text_from_file_bytes(file_bytes: bytes, filename: str) -> str:
    fn = filename.lower()
    try:
        if fn.endswith(".docx"):
            doc = Document(io.BytesIO(file_bytes))
            return "\n".join([p.text for p in doc.paragraphs]).strip()
        if fn.endswith(".pdf"):
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            return "\n".join([page.extract_text() or "" for page in reader.pages]).strip()
        if fn.endswith(".txt"):
            return file_bytes.decode("utf-8", errors="ignore")
    except:
        return ""
    return ""

def safe_parse_json(raw: str) -> dict:
    try: return json.loads(raw)
    except:
        m = re.search(r"\{.*\}", raw, re.DOTALL)
        if m:
            try: return json.loads(m.group(0))
            except: pass
        return {"title":"غير محدد","summary":raw[:1500],"keywords":[]}

def clean_filename(s: str) -> str:
    s = re.sub(r'[^\w\-\s]', '', (s or "").strip())
    s = s.replace(" ", "_")
    ts = datetime.datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    return f"{s[:60]}_{ts}"

def create_pdf_buffer(result: dict) -> io.BytesIO:
    buf = io.BytesIO()
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(name="Title", parent=styles["Title"], alignment=TA_CENTER)
    normal = styles["Normal"]
    doc = SimpleDocTemplate(buf, pagesize=A4)
    parts = []
    parts.append(Paragraph("🔍 تلخيص النص", title_style))
    parts.append(Spacer(1,6))
    parts.append(Paragraph(f"📘 العنوان: {result.get('title','')}", normal))
    parts.append(Spacer(1,6))
    parts.append(Paragraph("📝 الملخص:", normal))
    parts.append(Paragraph(result.get("summary",""), normal))
    if result.get("translation"):
        parts.append(Spacer(1,6))
        parts.append(Paragraph(f"🌍 الترجمة ({result.get('language','')}):", normal))
        parts.append(Paragraph(result.get("translation",""), normal))
    parts.append(Spacer(1,6))
    parts.append(Paragraph(f"🏷️ الكلمات المفتاحية: {', '.join(result.get('keywords',[]))}", normal))
    doc.build(parts)
    buf.seek(0)
    return buf

def create_docx_buffer(result: dict) -> io.BytesIO:
    buf = io.BytesIO()
    doc = Document()
    doc.add_heading("🔍 تلخيص النص", level=1)
    doc.add_paragraph(f"📘 العنوان: {result.get('title','')}")
    doc.add_paragraph("📝 الملخص:")
    doc.add_paragraph(result.get("summary",""))
    if result.get("translation"):
        doc.add_paragraph(f"🌍 الترجمة ({result.get('language','')}):")
        doc.add_paragraph(result.get("translation",""))
    doc.add_paragraph(f"🏷️ الكلمات المفتاحية: {', '.join(result.get('keywords',[]))}")
    doc.save(buf)
    buf.seek(0)
    return buf

def gemini_summarize(text: str) -> dict:
    model = genai.GenerativeModel("gemini-2.5-flash")
    prompt = f"""أعد المخرجات بصيغة JSON فقط:
{{"title":"...","summary":"...","keywords":["..."]}}
النص:
{text}
"""
    resp = model.generate_content(prompt)
    return safe_parse_json(resp.text)

def gemini_translate(text: str, to_lang: str) -> str:
    model = genai.GenerativeModel("gemini-2.5-flash")
    prompt = f"ترجم النص التالي إلى {to_lang} ترجمة دقيقة:\n\n{text}"
    resp = model.generate_content(prompt)
    return resp.text.strip()

# -------------------- Input --------------------
col_left, col_right = st.columns([2,1])
with col_left:
    tab_input = st.tabs(["✍️ إدخال نص","📁 رفع ملف"])
    with tab_input[0]:
        input_text = st.text_area("أدخل النص هنا:", height=280, placeholder="ألصق النص أو اكتب هنا...")
    with tab_input[1]:
        uploaded_file = st.file_uploader("ارفع ملف PDF / DOCX / TXT", type=["pdf","docx","txt"])
        input_text_from_file = ""
        if uploaded_file:
            bytes_data = uploaded_file.read()
            input_text_from_file = extract_text_from_file_bytes(bytes_data, uploaded_file.name)
            if not input_text_from_file:
                st.warning("تعذر استخراج نص من الملف.")
            else:
                st.success(f"تم استخراج النص (تقريبي طول: {len(input_text_from_file)} حرف).")
        else:
            input_text_from_file = ""

with col_right:
    st.markdown("### 🔧 الخيارات")
    target_lang = st.selectbox("ترجمة الملخص إلى:", ["No Translation","Arabic","English","French","Spanish"])
    action_btn = st.button("🚀 توليد الملخص")

final_input = input_text_from_file if uploaded_file else input_text or ""

# -------------------- Processing --------------------
if action_btn:
    if not final_input.strip():
        st.warning("يرجى إدخال نص أو رفع ملف صالح.")
    else:
        with st.spinner("⏳ جاري التلخيص..."):
            detected_lang = "unknown"
            try: detected_lang = detect(final_input)
            except: pass
            parsed = gemini_summarize(final_input)
            summary_text = parsed.get("summary","")
            translation_text = None
            final_lang = detected_lang
            if target_lang != "No Translation":
                translation_text = gemini_translate(summary_text, target_lang)
                final_lang = target_lang
            result = {
                "title": parsed.get("title","ملخص"),
                "summary": summary_text,
                "keywords": parsed.get("keywords",[]),
                "translation": translation_text,
                "language": final_lang,
                "input_text": final_input,
                "created_at": datetime.datetime.utcnow().isoformat()
            }
            # -------------------- Save to Supabase --------------------
            try:
                supabase.table("summaries").insert({
                    "title": result["title"],
                    "summary": result["summary"],
                    "translation": result.get("translation","
